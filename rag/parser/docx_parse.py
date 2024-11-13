import docx
class DocxParser:

    def __call__(self,file_path):
        file = docx.Document(file_path)
        content = ''
        # 遍历文档中的所有段落
        for paragraph in file.paragraphs:
            # 如果段落文本非空，则添加到内容字符串
            if paragraph.text.strip():
                content += paragraph.text + '\n'

        # 如果文档中包含表格，则遍历所有表格
        if len(file.tables) > 0:
            content += self.extract_tables(file.tables)

        # 返回文件内容而不是写入文件
        return content

    def extract_tables(self, tables, file_handle):
        for table in tables:
            max_cols = 0
            for row in table.rows:
                try:
                    # 获取实际存在的单元格数量
                    actual_cols = sum(1 for cell in row.cells if cell.text.strip())
                    max_cols = max(max_cols, actual_cols)
                except Exception as e:
                    print(f"Error processing row: {e}")
                    continue
            # 遍历表格中的每一行
            for row in table.rows:
                try:
                    # 获取实际存在的单元格
                    row_cells = [cell.text for cell in row.cells if cell.text.strip()]
                    # 填充缺失的单元格
                    row_cells += [''] * (max_cols - len(row_cells))
                    # 将一行中的所有单元格文本连接为一个字符串，用制表符分隔
                    row_str = '\t'.join(row_cells)
                    # 将处理好的行数据写入.txt文件
                    file_handle.write(row_str + '\n')
                except Exception as e:
                    print(f"Error processing row: {e}")
                    continue


